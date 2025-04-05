import argparse
from docx import Document
from re import search, split
from json import loads
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from extract import extract_all
from os import makedirs
from shutil import rmtree

def borrar_parrafo(parrafo) -> None:
    """
    Elimina un párrafo del documento de Word.

    @param parrafo: El párrafo que será eliminado del documento.
    """
    p = parrafo._element
    p.getparent().remove(p)
    p._p = p._element = None

def main(path, path_docx) -> None:
    """
    Procesa un archivo Jupyter Notebook (.ipynb), extrae las imágenes y el código generado,
    y los inserta en un documento de Word (.docx) siguiendo el formato IEEE.

    @param path: Ruta del archivo Jupyter Notebook (.ipynb) que contiene el código y salidas.
    @param path_docx: Ruta del archivo de Word (.docx) donde se insertarán las imágenes y el código.
    """
    # Extraemos el nombre base del archivo Jupyter Notebook (sin extensión)
    file_name = path.split('\\')[-1].split('.')[0]
    
    # Crear el directorio donde se guardarán las imágenes generadas (si no existe)
    makedirs(file_name, exist_ok=True)

    # Leer el archivo Jupyter Notebook como JSON
    with open(path, mode="r", encoding="utf-8") as f:
        file = loads(f.read())

        # Cargar el documento de Word donde se insertarán las imágenes y el código
        doc = Document(path_docx)
        flag = False

        # Llamada a la función para extraer imágenes y código del notebook
        extract_all(path, file_name)

        ## Eliminar párrafos que no son necesarios en el documento
        for p in doc.paragraphs:
            if search(r'^conclusiones$', p.text.lower()):
                break  # Detener al encontrar la sección de conclusiones
            if flag and not search(r'^código$', p.text.lower()):
                borrar_parrafo(p)  # Eliminar párrafos que no son parte del desarrollo o código
            if search(r'^desarrollo$', p.text.lower()):
                flag = True  # Iniciar la sección de desarrollo

        # Insertar la información y las imágenes del desarrollo
        for i, p in enumerate(doc.paragraphs):
            if search(r'^desarrollo$', p.text.lower()):
                c = 0  # Contador para las celdas del desarrollo
                for o, cell in enumerate(file['cells'][1:]):
                    if 'outputs' not in cell.keys():  # Si la celda no tiene salida (solo código)
                        if search(r'.*[0-9]\).*', cell['source'][0]):
                            e = split(r'[0-9]\)', cell['source'][0])[1].lstrip()
                            p = doc.paragraphs[i + c + 1].insert_paragraph_before(f"{e}", style='Heading 2')
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            c += 1
                    else:
                        # Si la celda tiene salida (imagen o gráfico generado)
                        if len(cell['outputs']) != 0:
                            p = doc.paragraphs[i + c + 1].insert_paragraph_before().add_run().add_picture(f"{file_name}/{o}img.png", width=Inches(3))
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            c += 1

        # Insertar las imágenes generadas del código
        for i, p in enumerate(doc.paragraphs):
            if search(r'^código$', p.text.lower()):
                c = 0  # Contador para las celdas de código
                for o, cell in enumerate(file['cells'][1:]):
                    if 'outputs' in cell.keys():  # Si la celda tiene salida (imagen del código)
                        p = doc.paragraphs[i + c + 1].insert_paragraph_before().add_run().add_picture(f"{file_name}/{o}code.png", width=Inches(3))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        c += 1
    
        # Guardar el documento con un nuevo nombre
        doc.save('new-'+path_docx)

    # Eliminar el directorio temporal creado para almacenar las imágenes
    rmtree(file_name)

if __name__ == '__main__':
    # Configurar el parser de argumentos para obtener los parámetros de entrada desde la línea de comandos
    parser = argparse.ArgumentParser(description="Procesar un archivo Jupyter Notebook e incrustar el código y las imágenes resultantes en un documento de Word previamente creado en formato IEEE.")
    parser.add_argument('path', type=str, help="Ruta del archivo Jupyter Notebook (.ipynb) de entrada.")
    parser.add_argument('path_docx', type=str, help="Ruta del archivo Word (.docx) de entrada donde se insertarán las imágenes y el código.")
    
    # Parsear los argumentos proporcionados desde la línea de comandos
    args = parser.parse_args()
    
    # Llamar a la función principal con los parámetros proporcionados
    main(args.path, args.path_docx)